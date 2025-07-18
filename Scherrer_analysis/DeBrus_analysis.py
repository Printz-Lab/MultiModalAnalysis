import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import cm
from scipy.integrate import quad
from scipy.optimize import root_scalar, curve_fit
from scipy.constants import hbar, pi, e, epsilon_0, m_e, h
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from docx import Document
from docx.shared import Inches
from pprint import pprint
import matplotlib as mpl

mpl.rcParams.update(
    {
        # 1) pick Arial for all sans-serif text…
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial"],
        # 2) make mathtext use Arial as well
        "mathtext.fontset": "custom",
        "mathtext.rm": "Arial",
        "mathtext.it": "Arial:italic",
        "mathtext.bf": "Arial:bold",
        "mathtext.default": "rm",
        # 3) still your other style settings
        "font.size": 14,
        "axes.labelsize": 18,
        "axes.titlesize": 18,
        "xtick.labelsize": 20,
        "ytick.labelsize": 20,
        "legend.fontsize": 14,
        "figure.figsize": (6, 8),
        "axes.linewidth": 1.5,
        "xtick.direction": "in",
        "ytick.direction": "in",
        "xtick.major.size": 6,
        "ytick.major.size": 6,
        "xtick.minor.size": 3,
        "ytick.minor.size": 3,
        "xtick.major.width": 1.2,
        "ytick.major.width": 1.2,
        "xtick.minor.width": 1.0,
        "ytick.minor.width": 1.0,
        "axes.grid": False,
        "savefig.dpi": 300,
        # if you had usetex on, turn it off so mathtext takes over:
        "text.usetex": False,
    }
)


# --- Physical models and robust root finding for Brus equation ---
def integrand(r, R, epsilon, N):
    """Function to integrate."""
    # sum over k from 1 to N
    sum_k = sum(
        ((r / R) ** (2 * k)) * (k + 1) / ((epsilon + 1) * k + 1)  # equation 11
        for k in range(1, N + 1)
    )
    return np.sin(pi * r / R) ** 2 * sum_k


def compute_integral(R, epsilon_pvk, epsilon_env, N=10):
    # Perform numerical integration from 0 to R
    integral_val, _ = quad(
        integrand,
        0,
        R,
        args=(epsilon_pvk, epsilon_env, N),  # note: leave out R, will fix below
    )
    # Actually quad passes R as first element of args, so we need correct call:
    # integral_val, _ = quad(integrand, 0, R, args=(R, epsilon_pvk, epsilon_env, N))
    return integral_val


def exciton_energy_eq(R, E_ex, me, mh, epsilon_pvk, epsilon_env, N=10):
    A = (h**2 / (8)) * (1 / me + 1 / mh)
    B = 1.8 * e**2 / (4 * pi * epsilon_pvk * epsilon_0)
    prefactor = (e**2 * (epsilon_pvk / epsilon_env - 1)) / (
        2 * pi * epsilon_pvk * epsilon_0
    )
    integral = compute_integral(R, epsilon_pvk, epsilon_env, N)
    C = prefactor * integral
    return A / R**2 - B / R - E_ex * e + C / R**2  # J


def solve_radius_robust(
    E_ex,
    me,
    mh,
    epsilon_pvk,
    epsilon_env,
    N=10,
    R_min=1e-10,
    R_max=2e-8,
    n_grid=50,
    tol=1e-12,
):
    """
    Find root R of exciton_energy_eq using a bracket search and Brent's method.
    Returns nan if no sign change is found or if endpoints are exact roots.
    """
    func = lambda R: exciton_energy_eq(R, E_ex, me, mh, epsilon_pvk, epsilon_env, N)

    # Safely evaluate endpoints
    try:
        f_min = func(R_min)
        f_max = func(R_max)
    except Exception:
        print(f"[ERROR] Exception evaluating endpoints: {R_min}, {R_max}")
        return np.nan

    # Check for endpoint zeros
    # if abs(f_min) < tol:
    #     print(f"[DEBUG] Found root at R_min: {R_min}")
    #     return R_min
    # if abs(f_max) < tol:
    #     print(f"[DEBUG] Found root at R_max: {R_max}")
    #     return R_max

    # Look for initial bracket
    if f_min * f_max > 0:
        # search grid for sign change
        Rs = np.logspace(np.log10(R_min), np.log10(R_max), n_grid)
        fs = np.array([func(r) for r in Rs])
        idx = np.where(fs[:-1] * fs[1:] < 0)[0]
        if idx.size == 0:
            return np.nan
        # use first sign-change bracket
        R_min, R_max = Rs[idx[0]], Rs[idx[0] + 1]
        f_min, f_max = func(R_min), func(R_max)

    # Ensure a valid bracket
    if f_min * f_max > 0:
        return np.nan

    # Use Brent's method on valid bracket
    try:
        sol = root_scalar(func, bracket=[R_min, R_max], method="brentq")
        return sol.root if sol.converged else np.nan
    except Exception:
        return np.nan


def exp_decay(t, A, tau, C):
    return A * np.exp(-t / tau) + C


def create_word_report(csv_path, df_fit, plot_paths):
    doc = Document()
    title = os.path.basename(csv_path)
    doc.add_heading(f"Report: {title}", level=1)
    doc.add_heading("Fit Parameters", level=2)
    table = doc.add_table(rows=1, cols=len(df_fit.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_fit.columns):
        hdr_cells[i].text = col
    for _, row in df_fit.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df_fit.columns):
            row_cells[i].text = f"{row[col]}"
    doc.add_heading("Plots", level=2)
    for plot in plot_paths:
        if os.path.exists(plot):
            doc.add_picture(plot, width=Inches(6))
    docx_path = os.path.splitext(csv_path)[0] + "_report.docx"
    doc.save(docx_path)
    return docx_path


def process_file(csv_path, target_time, regions):
    # Load and filter data
    df = pd.read_csv(csv_path)
    print(f"[INFO] Processing {csv_path}...")
    # print(df.head())
    time_col, peak_col, fwhm_col = df.columns[0], df.columns[3], df.columns[4]
    df_clean = df.dropna(subset=[peak_col])
    time = df_clean[time_col].to_numpy()[1:]
    peak_energy = df_clean[peak_col].to_numpy()[1:]
    fwhm = df_clean[fwhm_col].to_numpy()[1:]

    # Physical constants
    epsilon_pvk = 7.5  # dielectric constant of perovskite
    epsilon_env = 40  # dielectric constant of surrounding medium (e.g., vacuum)

    me_star = 0.138  # effective mass of electron in units of m_e
    mh_star = 0.118  # effective mass of hole in units of m_e
    me = me_star * m_e  # effective mass of electron in kg
    mh = mh_star * m_e  # effective mass of hole in kg

    # Determine bulk exciton energy at target_time
    idx_bulk = np.argmin(np.abs(time - target_time))
    Eg_bulk = peak_energy[idx_bulk]
    print(f"[INFO] Eg_bulk at t={time[idx_bulk]} s: {Eg_bulk:.4f} eV")

    radii = []
    for E in peak_energy:
        delta_E = E - Eg_bulk
        if delta_E <= 0:
            radii.append(np.nan)
            continue
        R = solve_radius_robust(delta_E, me, mh, epsilon_pvk, epsilon_env, N=20)
        radii.append(R * 1e9 if not np.isnan(R) else np.nan)

    if radii is None or len(radii) == 0:
        print(f"[ERROR] No valid radii found for {csv_path}.")
        return None, None, None
    radii = np.array(radii)
    # Filter noisy parts of radii (e.g., spikes or outliers)
    radii_filtered = radii.copy()
    if len(radii_filtered) > 5:
        # Use a rolling median filter to smooth out spikes
        window = 5
        radii_filtered = (
            pd.Series(radii_filtered)
            .rolling(window, center=True, min_periods=1)
            .median()
            .to_numpy()
        )
        # Optionally, mask outliers (e.g., >3 std from median)
        median = np.nanmedian(radii_filtered)
        std = np.nanstd(radii_filtered)
        outlier_mask = np.abs(radii_filtered - median) > 2.5 * std
        radii_filtered[outlier_mask] = np.nan
    radii = radii_filtered

    # Save analyzed CSV
    out_df = pd.DataFrame(
        {
            "Time (s)": time,
            "Peak Energy (eV)": peak_energy,
            "FWHM": fwhm,
            "Radius (nm)": radii,
        }
    )
    out_csv = os.path.splitext(csv_path)[0] + "_analysis.csv"
    out_df.to_csv(out_csv, index=False)

    # Plot 1: Energy & FWHM
    fig, ax1 = plt.subplots(figsize=(8, 6), dpi=300)
    ax1.plot(time, peak_energy, "s-", color="tab:blue", label="Peak Energy")
    ax1.set_ylabel("Peak Energy (eV)", color="tab:blue")
    ax1.set_xlabel("Time (s)")
    ax1.set_xlim(0, target_time)
    ax1.tick_params(axis="y", labelcolor="tab:blue")
    ax2 = ax1.twinx()
    ax2.plot(time, fwhm, "o-", color="tab:green", label="FWHM")
    ax2.set_ylabel("FWHM", color="tab:green")
    ax2.tick_params(axis="y", labelcolor="tab:green")
    plt.title(os.path.basename(csv_path))
    plt.tight_layout()
    fig.savefig(os.path.splitext(csv_path)[0] + "_energy_fwhm.png")
    # plt.show()
    plt.close(fig)

    # Plot 2: Energy & Radius with fits
    fit_params = []
    fig, ax1 = plt.subplots(figsize=(8, 6), dpi=300)
    ax1.plot(time, peak_energy, "s-", color="tab:blue", label="Peak Energy")
    ax1.set_ylabel("Peak Energy (eV)", color="tab:blue")
    ax1.set_xlabel("Time (s)")
    ax1.set_xlim(0, target_time)
    ax1.tick_params(axis="y", labelcolor="tab:blue")
    ax2 = ax1.twinx()
    ax2.plot(time, radii, "o-", color="tab:orange", label="Radius")
    for label, (t0, t1) in regions.items():
        mask = (time >= t0) & (time <= t1) & ~np.isnan(radii)
        t_seg, r_seg = time[mask], radii[mask]
        if len(t_seg) < 3:
            fit_params.append((label, np.nan, np.nan, np.nan, np.nan))
            continue
        p0 = [max(r_seg), (t1 - t0) / 2, min(r_seg)]
        try:
            popt, _ = curve_fit(exp_decay, t_seg, r_seg, p0=p0)
        except:
            popt = [np.nan, np.nan, np.nan]
        A, tau, C = popt
        mask = (time >= t0) & (time <= t1) & ~np.isnan(radii)
        ax2.plot(
            time[mask],
            exp_decay(time[mask], *popt),
            "--",
            label=f"{label} Fit",
            color="black",
        )
        ss_res = np.sum((r_seg - exp_decay(t_seg, *popt)) ** 2)
        ss_tot = np.sum((r_seg - np.mean(r_seg)) ** 2)
        R2 = 1 - ss_res / ss_tot if ss_tot > 0 else np.nan
        fit_params.append((label, A, tau, C, R2))
    ax2.set_ylabel("Radius (nm)", color="tab:orange")
    ax2.tick_params(axis="y", labelcolor="tab:orange")
    # print(radii)
    ax2.set_ylim(0, 40)
    plt.title(os.path.basename(csv_path))
    plt.tight_layout()
    fig.savefig(os.path.splitext(csv_path)[0] + "_energy_radius.png")
    # plt.show()
    plt.close(fig)

    # Plot 3: Growth rate
    fig, ax = plt.subplots(figsize=(8, 6), dpi=300)
    for label, A, tau, C, R2 in fit_params:
        t0, t1 = regions[label]
        mask = (time >= t0) & (time <= t1) & ~np.isnan(radii)
        ax.plot(time[mask], -A / tau * np.exp(-time[mask] / tau), "--", label=label)
    ax.set_xlabel("Time (s)")
    ax.set_ylabel("dRadius/dt (nm/s)")
    ax.legend()
    plt.tight_layout()
    fig.savefig(os.path.splitext(csv_path)[0] + "_growth_rate.png")
    # plt.show()
    plt.close(fig)
    plot1 = os.path.splitext(csv_path)[0] + "_energy_fwhm.png"
    plot2 = os.path.splitext(csv_path)[0] + "_energy_radius.png"
    plot3 = os.path.splitext(csv_path)[0] + "_growth_rate.png"
    # Save fit parameters
    df_fit = pd.DataFrame(fit_params, columns=["Region", "A", "Tau", "C", "R2"])
    df_fit.to_csv(os.path.splitext(csv_path)[0] + "_fit_params.csv", index=False)
    report = create_word_report(out_csv, df_fit, [plot1, plot2, plot3])
    return time, radii, report


def main():
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    # paths = []
    # while True:
    #     path = filedialog.askopenfilename(
    #         title="Select analysis CSVs", filetypes=[("CSV", "*.csv")]
    #     )
    #     if not path:
    #         break
    #     paths.append(path)
    #     if not messagebox.askyesno("Continue", "Add another file?"):
    #         break
    # if not paths:
    #     print("No files selected, exiting.")
    #     return
    # pprint(paths)
    paths = [
        "G:/MAPI_sean/MAPI_sean_control_S1_30_tube_5min/output/PL_FitResults.csv",
        "G:/MAPI_sean/MAPI_1pct_APA_S1_30_tube_5min/output/PL_FitResults.csv",
        "G:/MAPI_sean/MAPI_1pct_ABA_S1_30_tube_5min/output/PL_FitResults.csv",
        "G:/MAPI_sean/MAPI_1pct_AVA_S1_20_5min/output/PL_FitResults.csv",
        "G:/MAPI_sean/MAPI_1pct_AHA_S1_30_tube_5min/output/PL_FitResults.csv",
    ]
    labels = [
        "Pristine $MAPbI_3$",
        "1% 3-APA",
        "1% 4-ABA",
        "1% 5-AVA",
        "1% 7-AHA",
    ]
    colors_debrus = ["#080808", "#0262F3", "#863AB9", "#0A9628", "#E28812"]
    # colors_debrus = cm.tab10.colors[:len(labels)]
    markers = ["s", "o", "^", "v", "d"]
    target_time = 375  # end of experiment in seconds
    regions = {
        "Region1": (38, 84),
        "Region2": (84, 101),
        "Region3": (101, target_time - 1),
    }
    combined = []
    for p, label in zip(paths, labels):
        time, radii, report = process_file(p, target_time, regions)
        combined.append((label, time, radii))
        print(f"Processed {p}: report -> {report}")
    # combined plot
    plt.figure(figsize=(8, 6))
    for (label, time, radii), color, marker in zip(combined, colors_debrus, markers):
        mask = time > 0
        plt.plot(
            time[mask],
            radii[mask],
            # marker=marker,
            linestyle="-",
            label=label,
            color=color,
            alpha=0.8,
        )
    plt.xlabel("Time (s)")
    plt.ylabel("Radius (nm)")
    plt.legend()
    plt.tight_layout()
    out_comb = paths[0].replace(".csv", "_combined_plot.png")
    plt.savefig(out_comb, dpi=300)
    plt.show()
    print(f"Combined plot saved to {out_comb}")


if __name__ == "__main__":
    main()
